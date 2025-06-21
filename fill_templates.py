import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
import os
from datetime import datetime
import zipfile
import xml.etree.ElementTree as ET
import tempfile

# Importaciones para PDF
try:
    import win32com.client
    import pythoncom
    PDF_EXPORT_AVAILABLE = True
except ImportError:
    PDF_EXPORT_AVAILABLE = False

class TemplateFillerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Generador de Actas de Operatividad")
        self.root.geometry("800x600")
        self.style = ttk.Style()
        self.style.theme_use("clam")
        
        # Variable para guardar la ruta del último documento generado
        self.last_generated_docx = None
        
        # Configure styles
        self.style.configure("TFrame", background="#f0f0f0")
        self.style.configure("TLabel", background="#f0f0f0", font=("Segoe UI", 10))
        self.style.configure("TButton", font=("Segoe UI", 10), padding=6)
        self.style.configure("TCombobox", padding=5)
        self.style.configure("TEntry", padding=5)
        
        self.create_widgets()
    
    def create_widgets(self):
        # Main container
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Template selection
        ttk.Label(main_frame, text="Seleccionar Plantilla:").grid(row=0, column=0, sticky=tk.W, pady=(0, 10))
        self.template_var = tk.StringVar()
        self.template_cb = ttk.Combobox(main_frame, textvariable=self.template_var, values=["MOLDE 3G.docx", "MOLDE 4G.docx"], state="readonly")
        self.template_cb.grid(row=0, column=1, sticky=tk.EW, pady=(0, 10), padx=(10, 0))
        self.template_cb.bind("<<ComboboxSelected>>", self.on_template_select)
        
        # Fields frame
        self.fields_frame = ttk.Frame(main_frame)
        self.fields_frame.grid(row=1, column=0, columnspan=2, sticky=tk.NSEW)
        
        # Output filename
        ttk.Label(main_frame, text="Nombre del archivo:").grid(row=2, column=0, sticky=tk.W, pady=(20, 10))
        self.filename_var = tk.StringVar(value=f"ACTA_{datetime.now().strftime('%Y%m%d')}")
        self.filename_entry = ttk.Entry(main_frame, textvariable=self.filename_var)
        self.filename_entry.grid(row=2, column=1, sticky=tk.EW, pady=(20, 10), padx=(10, 0))
        
        # Output directory selection
        ttk.Label(main_frame, text="Carpeta de destino:").grid(row=3, column=0, sticky=tk.W, pady=(10, 10))
        self.output_dir_var = tk.StringVar(value=os.path.dirname(__file__))
        self.output_dir_frame = ttk.Frame(main_frame)
        self.output_dir_frame.grid(row=3, column=1, sticky=tk.EW, pady=(10, 10), padx=(10, 0))
        self.output_dir_entry = ttk.Entry(self.output_dir_frame, textvariable=self.output_dir_var, state='readonly')
        self.output_dir_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.browse_btn = ttk.Button(self.output_dir_frame, text="Examinar", command=self.browse_output_dir)
        self.browse_btn.pack(side=tk.RIGHT, padx=(5, 0))
        
        # Buttons frame
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.grid(row=4, column=0, columnspan=2, pady=(20, 0))
        
        # Generate button
        self.generate_btn = ttk.Button(buttons_frame, text="Generar DOCX", command=self.generate_document, state=tk.DISABLED)
        self.generate_btn.pack(side=tk.LEFT, padx=5)
        
        # PDF export button (only if available)
        if PDF_EXPORT_AVAILABLE:
            self.pdf_btn = ttk.Button(buttons_frame, text="Exportar a PDF", command=self.export_to_pdf, state=tk.DISABLED)
            self.pdf_btn.pack(side=tk.LEFT, padx=5)
            
            # Botón combinado: Generar DOCX y PDF
            self.combined_btn = ttk.Button(buttons_frame, text="Generar DOCX + PDF", command=self.generate_both, state=tk.DISABLED)
            self.combined_btn.pack(side=tk.LEFT, padx=5)
        else:
            no_pdf_label = ttk.Label(buttons_frame, text="(PDF export no disponible - Instalar pywin32)", foreground="red")
            no_pdf_label.pack(side=tk.LEFT, padx=5)
        
        # Status label
        self.status_label = ttk.Label(main_frame, text="Estado: Listo", foreground="green")
        self.status_label.grid(row=5, column=0, columnspan=2, pady=(10, 0))
        
        # Configure grid weights
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(1, weight=1)
        self.output_dir_frame.columnconfigure(0, weight=1)
    
    def on_template_select(self, event):
        # Clear previous fields
        for widget in self.fields_frame.winfo_children():
            widget.destroy()
        
        template = self.template_var.get()
        self.fields = {}
        
        if template == "MOLDE 3G.docx":
            self.create_3g_fields()
        elif template == "MOLDE 4G.docx":
            self.create_4g_fields()
        
        self.generate_btn.config(state=tk.NORMAL)
        if PDF_EXPORT_AVAILABLE:
            self.pdf_btn.config(state=tk.NORMAL)
            self.combined_btn.config(state=tk.NORMAL)
    
    def create_3g_fields(self):
        # Campos basados en el molde 3G real
        fields = [
            ("{{numero}}", "NUMERO"),
            ("{{empresa}}", "EMPRESA"),
            ("{{ruc}}", "RUC"),
            ("{{placa}}", "PLACA"),
            ("{{imei}}", "IMEI"),
            ("{{fec_ins}}", "FECHA DE INSTALACION"),
            ("{{dia}}", "DIA"),
            ("{{mes}}", "MES"),
        ]
        
        for i, (placeholder, label) in enumerate(fields):
            ttk.Label(self.fields_frame, text=f"{label}:").grid(row=i, column=0, sticky=tk.W, pady=5)
            entry = ttk.Entry(self.fields_frame, width=30)
            entry.grid(row=i, column=1, sticky=tk.EW, padx=(10, 0), pady=5)
            self.fields[placeholder] = entry
            
            # Auto-fill date fields
            if placeholder == "{{dia}}":
                entry.insert(0, str(datetime.now().day))
            elif placeholder == "{{mes}}":
                # Convertir número de mes a nombre
                meses = ["", "enero", "febrero", "marzo", "abril", "mayo", "junio",
                        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
                entry.insert(0, meses[datetime.now().month])
        
        self.fields_frame.columnconfigure(1, weight=1)
    
    def create_4g_fields(self):
        # Same fields for 4G template
        self.create_3g_fields()
    
    def browse_output_dir(self):
        """Permite al usuario seleccionar la carpeta de destino"""
        directory = filedialog.askdirectory(
            title="Selecciona la carpeta donde guardar los archivos",
            initialdir=self.output_dir_var.get()
        )
        if directory:
            self.output_dir_var.set(directory)
            self.update_status(f"Carpeta de destino: {os.path.basename(directory)}", "blue")
    
    def replace_in_xml_content(self, xml_content, replacements):
        """Reemplaza texto en contenido XML, incluyendo WordArt y elementos gráficos"""
        modified = False
        
        for placeholder, value in replacements.items():
            if placeholder in xml_content:
                xml_content = xml_content.replace(placeholder, str(value))
                modified = True
        
        return xml_content, modified
    
    def process_docx_with_xml(self, template_path, output_path, replacements):
        """Procesa el documento DOCX manipulando directamente el XML interno"""
        replacements_made = 0
        
        # Crear directorio temporal
        with tempfile.TemporaryDirectory() as temp_dir:
            # Extraer el DOCX (que es un ZIP)
            with zipfile.ZipFile(template_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # Lista de archivos XML que pueden contener texto
            xml_files_to_process = [
                'word/document.xml',
                'word/header1.xml',
                'word/header2.xml', 
                'word/header3.xml',
                'word/footer1.xml',
                'word/footer2.xml',
                'word/footer3.xml'
            ]
            
            # Procesar cada archivo XML
            for xml_file in xml_files_to_process:
                xml_path = os.path.join(temp_dir, xml_file)
                if os.path.exists(xml_path):
                    try:
                        # Leer el contenido XML
                        with open(xml_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                        
                        # Realizar reemplazos
                        new_content, modified = self.replace_in_xml_content(content, replacements)
                        
                        if modified:
                            # Escribir el contenido modificado
                            with open(xml_path, 'w', encoding='utf-8') as f:
                                f.write(new_content)
                            replacements_made += 1
                            print(f"Reemplazos realizados en: {xml_file}")
                    
                    except Exception as e:
                        print(f"Error procesando {xml_file}: {e}")
            
            # Buscar en archivos de relaciones y otros XMLs que puedan contener WordArt
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    if file.endswith('.xml') and file not in [f.split('/')[-1] for f in xml_files_to_process]:
                        file_path = os.path.join(root, file)
                        try:
                            with open(file_path, 'r', encoding='utf-8') as f:
                                content = f.read()
                            
                            new_content, modified = self.replace_in_xml_content(content, replacements)
                            
                            if modified:
                                with open(file_path, 'w', encoding='utf-8') as f:
                                    f.write(new_content)
                                replacements_made += 1
                                print(f"Reemplazos adicionales en: {file}")
                        
                        except Exception as e:
                            continue  # Ignorar errores en archivos no críticos
            
            # Recrear el DOCX
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arc_name = os.path.relpath(file_path, temp_dir)
                        zip_ref.write(file_path, arc_name)
        
        return replacements_made
    
    def replace_with_docx_library(self, doc, replacements):
        """Método tradicional usando python-docx como respaldo"""
        replaced_count = 0
        
        # Reemplazar en párrafos principales
        for paragraph in doc.paragraphs:
            if self.replace_in_paragraph(paragraph, replacements):
                replaced_count += 1
        
        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if self.replace_in_paragraph(paragraph, replacements):
                            replaced_count += 1
        
        # Reemplazar en headers y footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for paragraph in header.paragraphs:
                        if self.replace_in_paragraph(paragraph, replacements):
                            replaced_count += 1
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    if self.replace_in_paragraph(paragraph, replacements):
                                        replaced_count += 1
            
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for paragraph in footer.paragraphs:
                        if self.replace_in_paragraph(paragraph, replacements):
                            replaced_count += 1
                    for table in footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    if self.replace_in_paragraph(paragraph, replacements):
                                        replaced_count += 1
        
        return replaced_count
    
    def replace_in_paragraph(self, paragraph, replacements):
        """Reemplaza texto en un párrafo específico manteniendo el formato"""
        full_text = ""
        
        for run in paragraph.runs:
            full_text += run.text
        
        original_text = full_text
        for placeholder, value in replacements.items():
            full_text = full_text.replace(placeholder, str(value))
        
        if full_text != original_text:
            if paragraph.runs:
                first_run = paragraph.runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
                is_bold = first_run.font.bold
                is_italic = first_run.font.italic
                
                paragraph.clear()
                new_run = paragraph.add_run(full_text)
                
                try:
                    if font_name:
                        new_run.font.name = font_name
                    if font_size:
                        new_run.font.size = font_size
                    if is_bold:
                        new_run.font.bold = is_bold
                    if is_italic:
                        new_run.font.italic = is_italic
                except:
                    pass
            else:
                paragraph.text = full_text
            
            return True
        
        return False
    
    def update_status(self, message, color="black"):
        """Actualiza el estado en la interfaz"""
        self.status_label.config(text=f"Estado: {message}", foreground=color)
        self.root.update()
    
    def generate_document(self):
        """Genera solo el documento DOCX"""
        template = self.template_var.get()
        if not template:
            messagebox.showerror("Error", "Por favor selecciona una plantilla")
            return False
        
        output_name = self.filename_var.get().strip()
        if not output_name:
            output_name = f"ACTA_{datetime.now().strftime('%Y%m%d')}"
        
        if not output_name.endswith('.docx'):
            output_name += ".docx"
        
        # Preparar datos
        data = {}
        empty_fields = []
        
        for field, entry in self.fields.items():
            value = entry.get().strip()
            if not value:
                field_name = field.replace('{{', '').replace('}}', '').upper()
                empty_fields.append(field_name)
            data[field] = value
        
        if empty_fields:
            result = messagebox.askyesno("Campos vacíos", 
                                       f"Los siguientes campos están vacíos:\n{', '.join(empty_fields)}\n\n¿Deseas continuar?")
            if not result:
                return False
        
        try:
            self.update_status("Generando documento DOCX...", "blue")
            
            template_path = os.path.join(os.path.dirname(__file__), template)
            if not os.path.exists(template_path):
                messagebox.showerror("Error", f"No se encontró la plantilla: {template}")
                return False
            
            # Usar la carpeta seleccionada por el usuario
            output_dir = self.output_dir_var.get()
            if not os.path.exists(output_dir):
                messagebox.showerror("Error", f"La carpeta de destino no existe: {output_dir}")
                return False
            
            save_path = os.path.join(output_dir, output_name)
            
            print("Iniciando reemplazos...")
            print("Datos a reemplazar:")
            for key, value in data.items():
                print(f"  {key} -> {value}")
            
            # MÉTODO 1: Manipulación directa del XML (para WordArt y elementos gráficos)
            xml_replacements = self.process_docx_with_xml(template_path, save_path, data)
            print(f"Reemplazos XML realizados: {xml_replacements}")
            
            # MÉTODO 2: Usar python-docx como refuerzo (para elementos normales)
            doc = Document(save_path)
            docx_replacements = self.replace_with_docx_library(doc, data)
            print(f"Reemplazos python-docx realizados: {docx_replacements}")
            
            # Enforce Calibri font
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
            
            # Guardar cambios adicionales de python-docx
            if docx_replacements > 0:
                doc.save(save_path)
            
            # ¡IMPORTANTE! Guardar la ruta del documento generado
            self.last_generated_docx = save_path
            
            total_replacements = xml_replacements + docx_replacements
            
            if total_replacements > 0:
                self.update_status("Documento DOCX generado exitosamente", "green")
                messagebox.showinfo("Éxito", 
                                  f"✅ Documento generado exitosamente:\n\n"
                                  f"📁 Ubicación: {save_path}\n\n"
                                  f"📊 Total de reemplazos: {total_replacements}\n"
                                  f"   • XML (WordArt/Gráficos): {xml_replacements}\n"
                                  f"   • Texto normal: {docx_replacements}")
                
                # Preguntar si quiere abrir la carpeta
                result = messagebox.askyesno("Abrir carpeta", "¿Deseas abrir la carpeta donde se guardó el archivo?")
                if result:
                    import subprocess
                    subprocess.Popen(f'explorer "{os.path.dirname(save_path)}"')
            else:
                self.update_status("Documento generado con advertencias", "orange")
                messagebox.showwarning("Advertencia", 
                                     f"⚠️ Documento guardado en:\n\n"
                                     f"📁 {save_path}\n\n"
                                     f"No se detectaron reemplazos. Verifica los placeholders.")
            
            return True
        
        except Exception as e:
            self.update_status("Error al generar documento", "red")
            messagebox.showerror("Error", f"Error al generar el documento:\n{str(e)}")
            print(f"Error detallado: {e}")
            import traceback
            traceback.print_exc()
            return False

    def export_to_pdf(self):
        """Exporta el documento DOCX a PDF"""
        if not PDF_EXPORT_AVAILABLE:
            messagebox.showerror("Error", "Exportación a PDF no disponible.\nInstala pywin32: pip install pywin32")
            return
        
        # Verificar si existe un documento generado
        if not self.last_generated_docx or not os.path.exists(self.last_generated_docx):
            messagebox.showwarning("Advertencia", "Primero debes generar un documento DOCX")
            return
        
        try:
            self.update_status("Verificando Microsoft Word...", "blue")
            
            # Verificar que Word esté disponible
            pythoncom.CoInitialize()
            try:
                word = win32com.client.Dispatch("Word.Application")
            except Exception as e:
                raise Exception("Microsoft Word no está instalado o no es accesible")
            
            word.Visible = False
            
            self.update_status("Convirtiendo a PDF...", "blue")
            
            # Preparar nombre del archivo PDF
            output_name = self.filename_var.get().strip()
            if not output_name:
                output_name = f"ACTA_{datetime.now().strftime('%Y%m%d')}"
            
            # Remover extensión si existe y agregar .pdf
            if output_name.endswith('.docx'):
                output_name = output_name[:-5]
            output_name += ".pdf"
            
            # Usar la misma carpeta que el documento DOCX
            output_dir = self.output_dir_var.get()
            pdf_path = os.path.join(output_dir, output_name)
            
            print(f"Convirtiendo: {self.last_generated_docx} -> {pdf_path}")
            
            # Abrir documento y convertir
            doc = word.Documents.Open(os.path.abspath(self.last_generated_docx))
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = PDF format
            doc.Close()
            word.Quit()
            
            pythoncom.CoUninitialize()
            
            self.update_status("PDF generado exitosamente", "green")
            messagebox.showinfo("Éxito", f"✅ PDF generado exitosamente:\n\n📁 Ubicación: {pdf_path}")
            
            # Preguntamos si quiere abrir la carpeta
            result = messagebox.askyesno("Abrir carpeta", "¿Deseas abrir la carpeta donde se guardaron los archivos?")
            if result:
                import subprocess
                subprocess.Popen(f'explorer "{os.path.dirname(pdf_path)}"')
            
        except Exception as e:
            self.update_status("Error al generar PDF", "red")
            error_msg = str(e)
            if "Word" in error_msg:
                error_msg += "\n\nAsegúrate de que Microsoft Word esté instalado."
            messagebox.showerror("Error", f"Error al generar PDF:\n{error_msg}")
            print(f"Error detallado: {e}")
            import traceback
            traceback.print_exc()
    
    def generate_both(self):
        """Genera tanto el DOCX como el PDF"""
        if self.generate_document():
            self.export_to_pdf()

# Función para verificar dependencias
def check_dependencies():
    """Verifica las dependencias necesarias"""
    missing_deps = []
    
    try:
        import docx
    except ImportError:
        missing_deps.append("python-docx")
    
    try:
        import win32com.client
    except ImportError:
        missing_deps.append("pywin32 (para PDF)")
    
    if missing_deps:
        print("Dependencias faltantes:")
        for dep in missing_deps:
            print(f"  - {dep}")
        print("\nInstalar con:")
        print("pip install python-docx pywin32")

if __name__ == "__main__":
    check_dependencies()
    root = tk.Tk()
    app = TemplateFillerApp(root)
    root.mainloop()
